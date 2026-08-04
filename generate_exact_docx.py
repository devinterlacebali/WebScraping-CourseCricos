from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# Page setup - A4 portrait
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.0
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

def add_thin_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    t = '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>' if top else ''
    b = '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>' if bottom else ''
    l = '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>' if left else ''
    r = '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>' if right else ''
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>{t}{b}{l}{r}</w:tcBorders>')
    tcPr.append(tcBorders)

def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    tcPr.append(shading)

def set_cell_vertical_alignment(cell, align='center'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{align}"/>')
    tcPr.append(vAlign)

# =====================
# PAGE 1: CERTIFICATE
# =====================
doc.add_paragraph('')
doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CERTIFICATE')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
run.font.name = 'Calibri'

doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AWARDED TO')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
run.font.name = 'Calibri'

doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Anil Thapa')
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
run.font.name = 'Calibri'

doc.add_paragraph('')
doc.add_paragraph('')

# Qualification details box
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SHB30516')
run.font.size = Pt(16)
run.bold = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Certificate III in Barbering')
run.font.size = Pt(18)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
run.font.name = 'Calibri'

doc.add_paragraph('')
doc.add_paragraph('')

# Date & Certificate number
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Date of Completion:')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Calibri'
run = p.add_run('  24 March 2026')
run.font.size = Pt(12)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Certificate Number:')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Calibri'
run = p.add_run('  13493926-9353833')
run.font.size = Pt(12)
run.font.name = 'Calibri'

doc.add_paragraph('')
doc.add_paragraph('')

# Signature line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('_' * 40)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
run.font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Gareth Philpott')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Chief Executive Officer')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
run.font.name = 'Calibri'

doc.add_paragraph('')

# Footer line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('www.thebarberacademy.com.au | propellearning.com.au | RTO ID # 41209 | ABN 15 604 005 487')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Propel Education and Training Pty Ltd')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('This qualification is recognised within the Australian Qualifications Framework')
run.font.size = Pt(9)
run.italic = True
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = 'Calibri'

# PAGE BREAK
doc.add_page_break()

# =====================
# PAGE 2: RECORD OF RESULTS
# =====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('RECORD OF RESULTS')
run.font.size = Pt(20)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
run.font.name = 'Calibri'

doc.add_paragraph('')

# Student info
info_items = [
    ('NAME OF STUDENT:', 'Anil Thapa'),
    ('QUALIFICATION NAME:', 'Certificate III in Barbering'),
    ('QUALIFICATION CODE:', 'SHB30516'),
]
for label, val in info_items:
    p = doc.add_paragraph()
    run = p.add_run(label + '  ')
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run = p.add_run(val)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

doc.add_paragraph('')

# Units table
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set column widths
for row in table.rows:
    row.cells[0].width = Cm(3.2)
    row.cells[1].width = Cm(8.5)
    row.cells[2].width = Cm(1.8)
    row.cells[3].width = Cm(2.5)

hdr = table.rows[0].cells
headers = ['Unit of Competency Code', 'Unit of Competency Name', 'Outcome', 'Completion Date']
for i, h in enumerate(headers):
    hdr[i].text = ''
    p = hdr[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(8)
    run.font.name = 'Calibri'
    set_cell_shading(hdr[i], '1A3C6E')
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

units = [
    ('BSBSUS201', 'Participate in environmentally sustainable work practices', 'C', '25/02/2025'),
    ('SHBHBAS001', 'Provide shampoo and basin services', 'C', '05/08/2025'),
    ('SHBHCUT001', 'Design haircut structures', 'C', '11/03/2025'),
    ('SHBHCUT002', 'Create one length or solid haircut structures', 'C', '27/01/2026'),
    ('SHBHCUT003', 'Create graduated haircut structures', 'C', '05/08/2025'),
    ('SHBHCUT004', 'Create layered haircut structures', 'C', '10/02/2026'),
    ('SHBHCUT005', 'Cut hair using over-comb techniques', 'C', '30/10/2025'),
    ('SHBHCUT007', 'Create combined traditional and classic men\'s haircut structures', 'C', '10/02/2026'),
    ('SHBHCUT009', 'Cut hair using freehand clipper techniques', 'C', '04/02/2026'),
    ('SHBHCUT010', 'Create haircuts using tracks and carving', 'C', '13/02/2026'),
    ('SHBHCUT011', 'Design and maintain beards and moustaches', 'C', '06/02/2026'),
    ('SHBHCUT012', 'Shave heads and faces', 'C', '24/03/2026'),
    ('SHBHCUT013', 'Provide men\'s general grooming services', 'C', '03/11/2025'),
    ('SHBHDES001', 'Dry hair to shape', 'C', '03/11/2025'),
    ('SHBHIND001', 'Maintain and organise tools, equipment and work areas', 'C', '22/04/2025'),
    ('SHBHIND002', 'Research and use hairdressing industry information', 'C', '09/02/2026'),
    ('SHBHIND003', 'Develop and expand a client base', 'C', '22/10/2025'),
    ('SHBHTRI001', 'Identify and treat hair and scalp conditions', 'C', '05/08/2025'),
    ('SHBXCCS001', 'Conduct salon financial transactions', 'C', '16/03/2025'),
    ('SHBXCCS002', 'Provide salon services to clients', 'C', '05/08/2025'),
    ('SHBXIND001', 'Comply with organisational requirements within a personal services environment', 'C', '03/11/2025'),
    ('SHBXIND002', 'Communicate as part of a salon team', 'C', '05/08/2025'),
    ('SHBXWHS001', 'Apply safe hygiene, health and work practices', 'C', '25/02/2025'),
    ('SIRRINV001', 'Receive and handle retail stock', 'C', '10/04/2025'),
    ('SIRRMER001', 'Produce visual merchandise displays', 'C', '13/08/2025'),
    ('SIRXSLS001', 'Sell to the retail customer', 'C', '12/01/2026'),
]

even_row = False
for code, name, outcome, date in units:
    row = table.add_row()
    row.cells[0].width = Cm(3.2)
    row.cells[1].width = Cm(8.5)
    row.cells[2].width = Cm(1.8)
    row.cells[3].width = Cm(2.5)
    vals = [code, name, outcome, date]
    for i, val in enumerate(vals):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 2, 3] else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(val)
        run.font.size = Pt(8)
        run.font.name = 'Calibri'
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        if even_row:
            set_cell_shading(cell, 'F2F2F2')
        set_cell_vertical_alignment(cell, 'center')
    even_row = not even_row

doc.add_paragraph('')

# Signature
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('_' * 35)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
run.font.size = Pt(9)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('Gareth Philpott')
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('Chief Executive Officer')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('www.thebarberacademy.com.au | propellearning.com.au | RTO ID # 41209 | ABN 15 604 005 487')
run.font.size = Pt(7)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.font.name = 'Calibri'

# PAGE BREAK
doc.add_page_break()

# =====================
# PAGE 3: MASTER OF ACCOUNTING - TESTAMUR
# =====================
doc.add_paragraph('')
doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Master of Accounting')
run.font.size = Pt(24)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
run.font.name = 'Calibri'

doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('29 November 2024')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
run.font.name = 'Calibri'

doc.add_paragraph('')
doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Anil Thapa')
run.font.size = Pt(22)
run.bold = True
run.font.name = 'Calibri'

doc.add_paragraph('')
doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('24645')
run.font.size = Pt(14)
run.font.name = 'Calibri'

doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('138204_MBA_Testamur')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.font.name = 'Calibri'

doc.add_paragraph('')
doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Date Generated: 29 November 2024')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = 'Calibri'

# PAGE BREAK
doc.add_page_break()

# =====================
# PAGE 4-5: ACADEMIC TRANSCRIPT
# =====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('ANIL THAPA')
run.font.size = Pt(14)
run.bold = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
run = p.add_run('Student ID: ')
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Calibri'
run = p.add_run('1801398')
run.font.size = Pt(10)
run.font.name = 'Calibri'

p = doc.add_paragraph()
run = p.add_run('Birth date: ')
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Calibri'
run = p.add_run('30 March 1997')
run.font.size = Pt(10)
run.font.name = 'Calibri'

p = doc.add_paragraph()
run = p.add_run('26 November 2024')
run.font.size = Pt(10)
run.font.name = 'Calibri'

p = doc.add_paragraph()
run = p.add_run('Certificate number: ')
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Calibri'
run = p.add_run('24645')
run.font.size = Pt(10)
run.font.name = 'Calibri'

p = doc.add_paragraph()
run = p.add_run('Cumulative GPA: ')
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Calibri'
run = p.add_run('4.62')
run.font.size = Pt(10)
run.font.name = 'Calibri'

p = doc.add_paragraph()
run = p.add_run('Qualification: ')
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Calibri'
run = p.add_run('Master of Accounting')
run.font.size = Pt(10)
run.font.name = 'Calibri'

doc.add_paragraph('')

# Transcript table
t2 = doc.add_table(rows=1, cols=7)
t2.alignment = WD_TABLE_ALIGNMENT.CENTER

t2_hdr = t2.rows[0].cells
t2_headers = ['Year', 'Period', 'Subject', 'Subject Name', 'Mark', 'Grade', 'Credit\npoints']
for i, h in enumerate(t2_headers):
    t2_hdr[i].text = ''
    p = t2_hdr[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(8)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(t2_hdr[i], '1A3C6E')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

transcript = [
    ('2024', 'T2', 'ACCM4400', 'Auditing and Assurance', '70.00', 'C', '4'),
    ('2024', 'T2', 'ACCM6000', 'Capstone: Accounting and Governance', '61.00', 'P', '4'),
    ('2024', 'T2', 'FINM4000', 'Finance', '87.00', 'HD', '4'),
    ('2024', 'T1', 'ACCM4100', 'Management Accounting', '59.00', 'P', '4'),
    ('2024', 'T1', 'ACCM6000', 'Capstone: Accounting and Governance', '45.00', 'F', '0'),
    ('2024', 'T1', 'DATA4500', 'Social Media Analytics', '67.00', 'C', '4'),
    ('2023', 'T3', 'CLWM4100', 'Taxation Law', '53.00', 'P', '4'),
    ('2023', 'T3', 'ECOM4000', 'Economics', '58.00', 'P', '4'),
    ('2023', 'T2', 'ACCM4300', 'Financial Reporting', '77.00', 'D', '4'),
    ('2023', 'T2', 'DATA4900', 'Innovation and Creativity in Business Analytics', '68.00', 'C', '4'),
    ('2023', 'T2', 'FINM4100', 'Analytics in Accounting, Finance and Economics', '78.00', 'D', '4'),
    ('2023', 'T1', 'ACCM4200', 'Advanced Financial Accounting', '50.00', 'P', '4'),
]

even = False
for row_data in transcript:
    row = t2.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = ''
        p = row[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        run.font.size = Pt(8)
        run.font.name = 'Calibri'
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        if even:
            set_cell_shading(row[i], 'F2F2F2')
    even = not even

# "Continued over"
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('Continued over')
run.italic = True
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# PAGE BREAK
doc.add_page_break()

# Page 5 continued
t2_cont = doc.add_table(rows=1, cols=7)
t2_cont.alignment = WD_TABLE_ALIGNMENT.CENTER

t2c_hdr = t2_cont.rows[0].cells
for i, h in enumerate(t2_headers):
    t2c_hdr[i].text = ''
    p = t2c_hdr[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(8)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(t2c_hdr[i], '1A3C6E')

transcript_cont = [
    ('2023', 'T1', 'CLWM4000', 'Business and Corporations Law', '66.00', 'C', '4'),
    ('2023', 'T1', 'STAM4000', 'Quantitative Methods', '63.00', 'P', '4'),
    ('2022', 'T3', 'ACCM4000', 'Financial Accounting', '71.00', 'C', '4'),
    ('2022', 'T3', 'CISM4000', 'Information Systems in Accounting', '54.00', 'P', '4'),
    ('2022', 'T3', 'DATA4000', 'Introduction to Business Analytics', '65.00', 'C', '4'),
]

even = False
for row_data in transcript_cont:
    row = t2_cont.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = ''
        p = row[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        run.font.size = Pt(8)
        run.font.name = 'Calibri'
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        if even:
            set_cell_shading(row[i], 'F2F2F2')
    even = not even

doc.add_paragraph('')

# End of Record
p = doc.add_paragraph()
run = p.add_run('End of Record')
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Calibri'

p = doc.add_paragraph()
run = p.add_run('All Requirements for the Master of Accounting have been fulfilled')
run.font.size = Pt(10)
run.font.name = 'Calibri'

doc.add_paragraph('')

p = doc.add_paragraph()
run = p.add_run('Certified as a true and correct record')
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
run.font.name = 'Calibri'

doc.add_paragraph('')

p = doc.add_paragraph()
run = p.add_run('_' * 35)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
run.font.size = Pt(9)

p = doc.add_paragraph()
run = p.add_run('Rob Regan')
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Calibri'

p = doc.add_paragraph()
run = p.add_run('Principal Executive Officer')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
run.font.name = 'Calibri'

p = doc.add_paragraph()
run = p.add_run('Kaplan Business School')
run.font.size = Pt(10)
run.font.name = 'Calibri'

doc.add_paragraph('')

p = doc.add_paragraph()
run = p.add_run('The qualification is recognised within the Australian Qualifications Framework')
run.font.size = Pt(9)
run.italic = True
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = 'Calibri'

p = doc.add_paragraph()
run = p.add_run('Kaplan Business School is listed as an Institute of Higher Education on the Tertiary Education Quality and Standards Agency\'s National Register of Higher Education Providers, provider number (PRV12094).')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
run.font.name = 'Calibri'

p = doc.add_paragraph()
run = p.add_run('Kaplan Business School Pty Ltd ABN 86 098 181 947. CRICOS Provider Code 02426B.')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
run.font.name = 'Calibri'

# PAGE BREAK
doc.add_page_break()

# =====================
# PAGE 6: GRADING SYSTEM
# =====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Assessment Results')
run.font.size = Pt(18)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('Higher Education Grades')
run.font.size = Pt(12)
run.bold = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
run = p.add_run('Grades for subjects are reported as follows:')
run.font.size = Pt(10)
run.font.name = 'Calibri'

doc.add_paragraph('')

# Grading table
t3 = doc.add_table(rows=6, cols=4)
t3.alignment = WD_TABLE_ALIGNMENT.CENTER

grading = [
    ('Grade', 'Mark (%)', 'GPA', 'Comment'),
    ('High Distinction (HD)', '85-100', '7.0', ''),
    ('Distinction (D)', '75-84', '6.0', ''),
    ('Credit (C)', '65-74', '5.0', ''),
    ('Pass (P)', '50-64', '4.0', ''),
    ('Fail (F)', '0-49', '1.5', ''),
]
for i, row_data in enumerate(grading):
    for j, val in enumerate(row_data):
        cell = t3.cell(i, j)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        if i == 0:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(cell, '1A3C6E')
        set_cell_vertical_alignment(cell, 'center')

doc.add_paragraph('')
doc.add_paragraph('')

# Additional grade info
grade_info = [
    ('Withdrawn Fail (WF)', '0.0', 'Indicates a student has formally notified Kaplan Business School of their withdrawal from a subject after the census date and prior to the final day of teaching in that trimester.'),
    ('Absent Fail (AF)', '0.0', 'Indicates that a student did not submit or sit any assessment events for a subject and the student did not formally withdraw from the subject.'),
    ('Withdrawn No Fail (WNF)', '', 'Indicates that a student has formally notified Kaplan Business School of their withdrawal from the subject prior to the census date.'),
    ('Incomplete (I)', '', 'Indicates that a student has not had a final grade determined because (i) they have not completed all assessment tasks and therefore required an extension of time, or (ii) they have been granted a supplementary examination or additional assessment item. The grade must be finalised before the end of the following trimester.'),
    ('Exempt (E)', '', 'Indicates that a student has achieved an exemption for the subject via the Recognition of Prior Learning process.'),
]

for grade_name, gpa, desc in grade_info:
    p = doc.add_paragraph()
    run = p.add_run(grade_name)
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Calibri'
    if gpa:
        run = p.add_run(f'  ({gpa})')
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
    
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.5)
    run = p2.add_run(desc)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = 'Calibri'

doc.add_paragraph('')

p = doc.add_paragraph()
run = p.add_run('Please refer to the Assessment Policy available at www.kaplan.edu.au for more information regarding assessment results.')
run.font.size = Pt(9)
run.italic = True
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = 'Calibri'

doc.add_paragraph('')

p = doc.add_paragraph()
run = p.add_run('Grade Point Average (GPA): ')
run.bold = True
run.font.size = Pt(9)
run.font.name = 'Calibri'
run = p.add_run('The average result of all the grades achieved throughout the course calculated on a 7-point grading scale where 7 is the highest and 0 is the lowest achievement. Fail grades are included in the calculation.')
run.font.size = Pt(9)
run.font.name = 'Calibri'

doc.add_paragraph('')

p = doc.add_paragraph()
run = p.add_run('Specialisation: ')
run.bold = True
run.font.size = Pt(9)
run.font.name = 'Calibri'
run = p.add_run('To graduate with a specialisation in Kaplan Business School\'s Master of Business Administration course, 3 of the 12 completed subjects must relate directly to the area of specialisation.')
run.font.size = Pt(9)
run.font.name = 'Calibri'

# Save
output = r'C:\Users\Dewa(Interlace)\Documents\Interlace Code\WebScraping-CourseCricos\Onshore_Academics_Anil_Thapa.docx'
doc.save(output)
print(f'DOCX saved: {output}')
