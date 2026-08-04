from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.0
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    tcPr.append(shading)

def set_cell_valign(cell, align='center'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{align}"/>')
    tcPr.append(vAlign)

def make_header_row(table, headers, widths=None):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1A3C6E')
        set_cell_valign(cell, 'center')
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

def add_data_row(table, values, widths=None, even=False):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(str(val))
        run.font.size = Pt(8)
        run.font.name = 'Calibri'
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        if even:
            set_cell_shading(cell, 'F2F2F2')
        set_cell_valign(cell, 'center')

def add_heading_line(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    return h

def add_info(label, value):
    p = doc.add_paragraph()
    r = p.add_run(label + '  ')
    r.bold = True
    r.font.size = Pt(10)
    r.font.name = 'Calibri'
    r2 = p.add_run(value)
    r2.font.size = Pt(10)
    r2.font.name = 'Calibri'
    return p

# =========================================
# PAGE 1: TITLE
# =========================================
doc.add_paragraph('')
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('OFFSHORE ACADEMICS')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Anil Thapa')
run.font.size = Pt(22)
run.bold = True
run.font.name = 'Calibri'

doc.add_paragraph('')

info_table = doc.add_table(rows=6, cols=2)
info_table.style = 'Light Grid Accent 1'
for i, (label, val) in enumerate([
    ('Full Name', 'Anil Thapa'),
    ('Date of Birth', '30 March 1997 (AD) / 2053/12/17 (BS)'),
    ('Father\'s Name', 'Surya Bahadur Thapa'),
    ('Mother\'s Name', 'Sunita Thapa'),
    ('Permanent Address', 'Nilkantha-5, Dhading, Nepal'),
    ('Total Offshore Qualifications', '3 (BBA + Grade XII + Grade XI)'),
]):
    info_table.cell(i, 0).text = label
    info_table.cell(i, 1).text = val
    for cell in info_table.rows[i].cells:
        for p in cell.paragraphs:
            p.style.font.size = Pt(10)

doc.add_paragraph('')

# =========================================
# SECTION 1: BBA from Pokhara University
# =========================================
add_heading_line('1. Bachelor of Business Administration (BBA)', level=1)

add_info('University:', 'Pokhara University, Office of the Controller of Examinations, Kaski, Nepal')
add_info('College:', 'Pokhara College of Management')
add_info('Program:', 'Bachelor of Business Administration')
add_info('Level:', 'Bachelor')
add_info('Date of Registration:', 'September 2016')
add_info('Date of Completion:', 'June 2021')
add_info('Transcript Serial:', '67891')
add_info('Exam Roll No.:', '17030259')
add_info('Registration No.:', '2016-2-03-0149')
add_info('Total Credits:', '120.00')
add_info('CGPA:', '2.52 / 4.0')

doc.add_paragraph('')
add_heading_line('Semester-wise Academic Record', level=2)

# Semester 1 & 2
p = doc.add_paragraph()
r = p.add_run('First & Second Semester')
r.bold = True
r.font.size = Pt(10)
r.font.name = 'Calibri'

t1 = doc.add_table(rows=1, cols=10)
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
make_header_row(t1, ['SN', 'Code', 'Course Title', 'Cr', 'Grade', 'SN', 'Code', 'Course Title', 'Cr', 'Grade'])

sem12 = [
    ('1', 'ENG 101', 'English I', '3', 'C+', '1', 'ENG 102', 'English II', '3', 'C-'),
    ('2', 'MTH 101', 'Business Mathematics I', '3', 'B-', '2', 'MTH 102', 'Business Mathematics II', '3', 'C'),
    ('3', 'ACC 121', 'Financial Accounting I', '3', 'B-', '3', 'ACC 122', 'Financial Accounting II', '3', 'C'),
    ('4', 'MGT 111', 'Principles of Management', '3', 'C', '4', 'PSY 101', 'General Psychology', '3', 'C+'),
    ('5', 'MIS 101', 'Computer and IT Applications', '3', 'C', '5', 'ECO 101', 'Introductory Microeconomics', '3', 'C'),
]
even = False
for vals in sem12:
    add_data_row(t1, vals, even=even)
    even = not even

# Totals row
row = t1.add_row()
for i, val in enumerate(['', '', 'Total', '15', '', '', '', 'Total', '15', '']):
    cell = row.cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(val)
    run.bold = True
    run.font.size = Pt(8)
    run.font.name = 'Calibri'
    set_cell_shading(cell, 'E8E8E8')

row = t1.add_row()
for i, val in enumerate(['', '', 'SGPA', '2.34', '', '', '', 'SGPA', '2.00', '']):
    cell = row.cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(val)
    run.bold = True
    run.font.size = Pt(8)
    run.font.name = 'Calibri'
    set_cell_shading(cell, 'E8E8E8')

doc.add_paragraph('')

# Semester 3 & 4
p = doc.add_paragraph()
r = p.add_run('Third & Fourth Semester')
r.bold = True
r.font.size = Pt(10)
r.font.name = 'Calibri'

t2 = doc.add_table(rows=1, cols=10)
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
make_header_row(t2, ['SN', 'Code', 'Course Title', 'Cr', 'Grade', 'SN', 'Code', 'Course Title', 'Cr', 'Grade'])

sem34 = [
    ('1', 'ENG 201', 'Business Communication I', '3', 'C-', '1', 'ENG 202', 'Business Communication II', '3', 'C'),
    ('2', 'STT 101', 'Business Statistics', '3', 'B', '2', 'STT 201', 'Data Analysis and Modeling', '3', 'B'),
    ('3', 'FIN 131', 'Essentials of Finance', '3', 'B', '3', 'MGT 211', 'Fundamentals of Organisational Behaviour', '3', 'C'),
    ('4', 'SOC 101', 'Fundamentals of Sociology', '3', 'C', '4', 'MKT 241', 'Principles of Marketing', '3', 'C'),
    ('5', 'ECO 201', 'Introductory Macroeconomics', '3', 'C+', '5', 'FIN 231', 'Financial Management', '3', 'B'),
]
even = False
for vals in sem34:
    add_data_row(t2, vals, even=even)
    even = not even

row = t2.add_row()
for i, val in enumerate(['', '', 'Total', '15', '', '', '', 'Total', '15', '']):
    cell = row.cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(val)
    run.bold = True
    run.font.size = Pt(8)
    run.font.name = 'Calibri'
    set_cell_shading(cell, 'E8E8E8')

row = t2.add_row()
for i, val in enumerate(['', '', 'SGPA', '2.34', '', '', '', 'SGPA', '2.40', '']):
    cell = row.cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(val)
    run.bold = True
    run.font.size = Pt(8)
    run.font.name = 'Calibri'
    set_cell_shading(cell, 'E8E8E8')

doc.add_paragraph('')

# PAGE BREAK - Semester 5-8
doc.add_page_break()

p = doc.add_paragraph()
r = p.add_run('Fifth & Sixth Semester')
r.bold = True
r.font.size = Pt(10)
r.font.name = 'Calibri'

t3 = doc.add_table(rows=1, cols=10)
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
make_header_row(t3, ['SN', 'Code', 'Course Title', 'Cr', 'Grade', 'SN', 'Code', 'Course Title', 'Cr', 'Grade'])

sem56 = [
    ('1', 'ACC 221', 'Basics of Managerial Accounting', '3', 'C', '1', 'MIS 201', 'Intro to Management Information Systems', '3', 'C'),
    ('2', 'MGT 311', 'Fundamentals of Operations Management', '3', 'C+', '2', 'LAW 291', 'Legal Aspects of Business and Technology', '3', 'B-'),
    ('3', 'RCH 311', 'Business Research Methods', '3', 'B-', '3', 'MGT 212', 'Business and Society', '3', 'C+'),
    ('4', 'MGT 314', 'Management of Human Resources', '3', 'C-', '4', 'PRJ 491', 'Project Work', '3', 'A-'),
    ('5', 'MKT 442', '*Advertising and Sales Promotion', '3', 'B', '5', 'MKT 443', '*Sales Management', '3', ''),
    ('6', 'FIN 437', '*Financial Institutions and Markets', '3', 'B', '6', 'FIN 433', '*Investment Management', '3', 'B-'),
]
even = False
for vals in sem56:
    add_data_row(t3, vals, even=even)
    even = not even

row = t3.add_row()
for i, val in enumerate(['', '', 'Total', '15', '', '', '', 'Total', '15', '']):
    cell = row.cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(val)
    run.bold = True; run.font.size = Pt(8); run.font.name = 'Calibri'
    set_cell_shading(cell, 'E8E8E8')

row = t3.add_row()
for i, val in enumerate(['', '', 'SGPA', '2.34', '', '', '', 'SGPA', '2.68', '']):
    cell = row.cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(val)
    run.bold = True; run.font.size = Pt(8); run.font.name = 'Calibri'
    set_cell_shading(cell, 'E8E8E8')

doc.add_paragraph('')

p = doc.add_paragraph()
r = p.add_run('Seventh & Eighth Semester')
r.bold = True
r.font.size = Pt(10)
r.font.name = 'Calibri'

t4 = doc.add_table(rows=1, cols=10)
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
make_header_row(t4, ['SN', 'Code', 'Course Title', 'Cr', 'Grade', 'SN', 'Code', 'Course Title', 'Cr', 'Grade'])

sem78 = [
    ('1', 'MGT 411', 'Business Environment in Nepal', '3', 'B-', '1', 'MGT 412', 'Strategic Management', '3', 'C+'),
    ('2', 'MGT 312', 'Fundamentals of Entrepreneurship', '3', 'B-', '2', 'MGT 313', 'Introduction to International Business', '3', 'B-'),
    ('3', 'INT 391', 'Internship', '3', 'A-', '3', 'MIS 301', 'Essentials of e-Business', '3', 'B'),
    ('4', 'NBE 396', '*Management of Technology', '3', 'C+', '4', 'NBE 394', '*Media and Public Relations', '3', 'B+'),
    ('5', 'FIN 431', '*Corporate Finance', '3', 'A-', '5', 'FIN 434', '*Bank Operations and Management', '3', 'A-'),
]
even = False
for vals in sem78:
    add_data_row(t4, vals, even=even)
    even = not even

row = t4.add_row()
for i, val in enumerate(['', '', 'Total', '15', '', '', '', 'Total', '15', '']):
    cell = row.cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(val)
    run.bold = True; run.font.size = Pt(8); run.font.name = 'Calibri'
    set_cell_shading(cell, 'E8E8E8')

row = t4.add_row()
for i, val in enumerate(['', '', 'SGPA', '3.02', '', '', '', 'SGPA', '3.00', '']):
    cell = row.cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(val)
    run.bold = True; run.font.size = Pt(8); run.font.name = 'Calibri'
    set_cell_shading(cell, 'E8E8E8')

doc.add_paragraph('')

p = doc.add_paragraph()
r = p.add_run('Summary: ')
r.bold = True; r.font.size = Pt(10)
r2 = p.add_run('Total Credits: 120.00  |  CGPA: 2.52')
r2.font.size = Pt(10)

doc.add_paragraph('')

# =========================================
# SECTION 2: Pokhara University Certificates
# =========================================
add_heading_line('2. Pokhara University — Certificates', level=1)

add_heading_line('2A. Provisional Certificate', level=2)
add_info('Issue Date:', '04 February 2022')
add_info('Issue No.:', '48975')
add_info('CGPA:', '2.52')
add_info('Program:', 'Bachelor of Business Administration')
add_info('Period:', '2016 to 2021')
add_info('Institution:', 'Pokhara College of Management')

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
p.paragraph_format.space_before = Pt(6)
run = p.add_run('"This is to certify that Mr. ANIL THAPA was a student of POKHARA COLLEGE OF MANAGEMENT from 2016 to 2021. S/he has completed all the requirements of the Bachelor of Business Administration Program with a CGPA score of 2.52."')
run.font.size = Pt(9)
run.italic = True
run.font.name = 'Calibri'

add_heading_line('2B. Migration Certificate', level=2)
add_info('Issue Date:', '04 February 2022')
add_info('Registration No.:', '2016-2-03-0149')
add_info('Red Stamp No.:', '72922')

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
p.paragraph_format.space_before = Pt(6)
run = p.add_run('"This is to certify that Mr. ANIL THAPA, son of SURYA BAHADUR THAPA, of POKHARA COLLEGE OF MANAGEMENT, was a student of this University. Pokhara University has no objection in his/her joining any other University or Institution."')
run.font.size = Pt(9)
run.italic = True
run.font.name = 'Calibri'

doc.add_page_break()

# =========================================
# SECTION 3: NEB Grade XI & XII
# =========================================
add_heading_line('3. National Examinations Board (NEB) — Grade XI & XII', level=1)

add_info('Student:', 'Anil Thapa')
add_info('Date of Birth:', '2053/12/17 (BS) = 30 March 1997 (AD)')
add_info('School:', 'Motherland Secondary School, Masbar, Pokhara-7, Kaski')
add_info('NEB Registration No.:', '714036022')
add_info('Serial No.:', '06-001724')
add_info('Transcript No.:', 'A0021926')

doc.add_paragraph('')
add_heading_line('Grade XI (Year 2071 BS / 2014 AD)', level=2)

t5 = doc.add_table(rows=1, cols=5)
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
make_header_row(t5, ['Subject', 'Full Marks', 'Pass Marks', 'Marks Secured', 'Remarks'])

gr11 = [
    ('COMPULSORY ENGLISH', '100', '35', '51', ''),
    ('COMPULSORY NEPALI', '100', '35', '60', ''),
    ('ACCOUNTANCY', '100', '35', '74', ''),
    ('ECONOMICS', '100', '35', '63', ''),
    ('COMPUTER SCIENCE (TH)', '075', '27', '44', ''),
    ('COMPUTER SCIENCE (PR)', '025', '10', '23', ''),
]
even = False
for vals in gr11:
    add_data_row(t5, vals, even=even)
    even = not even

row = t5.add_row()
for i, val in enumerate(['Total', '500', '', '315', '']):
    cell = row.cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(val)
    run.bold = True; run.font.size = Pt(8); run.font.name = 'Calibri'
    set_cell_shading(cell, 'E8E8E8')

add_info('Symbol Number:', '14005023')

doc.add_paragraph('')

add_heading_line('Grade XII (Year 2072 BS / 2015 AD)', level=2)

t6 = doc.add_table(rows=1, cols=5)
t6.alignment = WD_TABLE_ALIGNMENT.CENTER
make_header_row(t6, ['Subject', 'Full Marks', 'Pass Marks', 'Marks Secured', 'Remarks'])

gr12 = [
    ('COMPULSORY ENGLISH', '100', '35', '57', ''),
    ('ACCOUNTANCY', '100', '35', '70', ''),
    ('ECONOMICS', '100', '35', '52', ''),
    ('COMPUTER SCIENCE (TH)', '075', '27', '43', ''),
    ('COMPUTER SCIENCE (PR)', '025', '10', '23', ''),
    ('BUSINESS MATHEMATICS', '100', '35', '69', '*'),
]
even = False
for vals in gr12:
    add_data_row(t6, vals, even=even)
    even = not even

row = t6.add_row()
for i, val in enumerate(['Total', '500', '', '314', '']):
    cell = row.cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(val)
    run.bold = True; run.font.size = Pt(8); run.font.name = 'Calibri'
    set_cell_shading(cell, 'E8E8E8')

add_info('Symbol Number:', '24004544 / 90602250')

doc.add_paragraph('')

# Overall result
add_heading_line('Overall Result', level=2)

t7 = doc.add_table(rows=4, cols=2)
t7.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, val) in enumerate([
    ('Grand Total', '629 / 1000'),
    ('Percentage', '62.90%'),
    ('Division', 'First Division'),
    ('Year of Completion', '2072 BS (2015 AD)'),
]):
    t7.cell(i, 0).text = label
    t7.cell(i, 1).text = val
    if i == 0:
        for cell in t7.rows[i].cells:
            set_cell_shading(cell, '1A3C6E')
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.bold = True
                    r.font.size = Pt(10)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

# Grading System
add_heading_line('NEB Grading System', level=2)

grading_data = [
    ('75% above', 'Distinction'),
    ('60% above', 'First Division'),
    ('35% above', 'Pass Division'),
]
t8 = doc.add_table(rows=1, cols=2)
t8.alignment = WD_TABLE_ALIGNMENT.CENTER
make_header_row(t8, ['Marks Range', 'Division'])
for val in grading_data:
    add_data_row(t8, val)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
run = p.add_run('* Means a student has passed in the second attempt.')
run.font.size = Pt(9); run.italic = True

doc.add_page_break()

# =========================================
# SECTION 4: Transfer/Character Certificate
# =========================================
add_heading_line('4. Transfer / Character Certificate', level=1)

add_info('School:', 'Motherland Higher Secondary School, Masbar, Pokhara-7')
add_info('Established:', '2051 BS')
add_info('Registration No.:', '714036022')
add_info('Symbol No.:', '90602250')
add_info('Date Issued:', '2073/09/24 BS = 2017/01/08 AD')

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
p.paragraph_format.space_before = Pt(6)
run = p.add_run('"Dedicated to: Quality education, where learning is infinite"')
run.font.size = Pt(9); run.italic = True; run.font.name = 'Calibri'

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
p.paragraph_format.space_before = Pt(6)
run = p.add_run('"This is to certify that Anil Thapa, the son of Mr. Surya Bahadur Thapa and Mrs. Sunita Thapa, an inhabitant of Nilkantha-5, Dhading, Nepal, was a bona fide student of this institution. He passed the Final examinations of classes 11 and 12 in Management Stream conducted by National Examination Board, Nepal, in 2072 B.S. (2015 A.D.) and was placed in First Division. He bears a good moral character."')
run.font.size = Pt(9); run.italic = True; run.font.name = 'Calibri'

add_info('Contact:', 'PO Box 185, Phone: 977-61-463256/464256/465377/461556')
add_info('Website:', 'www.motherland.edu.np')

doc.add_paragraph('')

# =========================================
# SECTION 5: NEB Provisional Certificate
# =========================================
add_heading_line('5. NEB Provisional Certificate', level=1)

add_info('Serial No.:', 'P0021927')
add_info('NEB Registration No.:', '714036022')
add_info('Symbol No.:', '90602250')
add_info('Student:', 'ANIL THAPA')
add_info('School:', 'Motherland Secondary School, Kaski')
add_info('Examination Year:', '2072 BS (2015 AD)')
add_info('Division:', 'First Division')
add_info('Date of Birth:', '2053/12/17 BS')
add_info('Date of Issue:', '2023/08/23')
add_info('Location:', 'Sanothimi, Bhaktapur')

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
p.paragraph_format.space_before = Pt(6)
run = p.add_run('"The original certificate is not yet issued to him/her."')
run.font.size = Pt(9); run.italic = True; run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00); run.font.name = 'Calibri'

doc.add_paragraph('')

# =========================================
# FINAL: Summary
# =========================================
add_heading_line('Summary', level=1)

p = doc.add_paragraph()
p.add_run('Anil Thapa holds the following offshore qualifications:\n')
bullets = [
    'Bachelor of Business Administration (BBA) — Pokhara University, Nepal — Completed June 2021 — CGPA 2.52/4.0',
    'Grade XII (Higher Secondary) — National Examinations Board, Nepal — Completed 2072 BS (2015 AD) — First Division (62.90%)',
    'Grade XI — National Examinations Board, Nepal — Completed 2071 BS (2014 AD) — 315/500 marks',
    'Transfer/Character Certificate — Motherland Higher Secondary School — Issued January 2017',
    'Migration Certificate — Pokhara University — Issued February 2022',
    'Provisional Certificate — Pokhara University — Issued February 2022',
    'NEB Provisional Certificate — National Examinations Board — Issued August 2023',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# Save
output = r'C:\Users\Dewa(Interlace)\Documents\Interlace Code\WebScraping-CourseCricos\Offshore_Academics_Anil_Thapa.docx'
doc.save(output)
print(f'DOCX saved: {output}')
