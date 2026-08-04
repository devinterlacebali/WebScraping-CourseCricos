from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# TITLE
title = doc.add_heading('ONSHORE ACADEMICS — Anil Thapa', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

# Info table
info_table = doc.add_table(rows=4, cols=2)
info_table.style = 'Light Grid Accent 1'
for i, (label, val) in enumerate([
    ('Full Name', 'Anil Thapa'),
    ('Date of Birth', '30 March 1997'),
    ('Student ID (Kaplan)', '1801398'),
    ('Certificate Number', '13493926-9353833'),
]):
    info_table.cell(i, 0).text = label
    info_table.cell(i, 1).text = val
doc.add_paragraph('')

# ===== SECTION 1 =====
doc.add_heading('1. Certificate III in Barbering', level=1)

for label, val in [
    ('Qualification:', 'SHB30516 — Certificate III in Barbering'),
    ('Date of Completion:', '24 March 2026'),
    ('Training Provider:', 'Propel Education and Training Pty Ltd (RTO #41209)'),
    ('Website:', 'www.thebarberacademy.com.au | propellearning.com.au'),
    ('Chief Executive:', 'Gareth Philpott'),
]:
    p = doc.add_paragraph()
    r = p.add_run(label + ' ')
    r.bold = True
    p.add_run(val)

doc.add_heading('Units Completed (26 units — All Competent)', level=2)

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['Unit Code', 'Unit Name', 'Completion Date']):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

units = [
    ('BSBSUS201', 'Participate in environmentally sustainable work practices', '25/02/2025'),
    ('SHBHBAS001', 'Provide shampoo and basin services', '05/08/2025'),
    ('SHBHCUT001', 'Design haircut structures', '11/03/2025'),
    ('SHBHCUT002', 'Create one length or solid haircut structures', '27/01/2026'),
    ('SHBHCUT003', 'Create graduated haircut structures', '05/08/2025'),
    ('SHBHCUT004', 'Create layered haircut structures', '10/02/2026'),
    ('SHBHCUT005', 'Cut hair using over-comb techniques', '30/10/2025'),
    ('SHBHCUT007', 'Create combined traditional and classic mens haircut structures', '10/02/2026'),
    ('SHBHCUT009', 'Cut hair using freehand clipper techniques', '04/02/2026'),
    ('SHBHCUT010', 'Create haircuts using tracks and carving', '13/02/2026'),
    ('SHBHCUT011', 'Design and maintain beards and moustaches', '06/02/2026'),
    ('SHBHCUT012', 'Shave heads and faces', '24/03/2026'),
    ('SHBHCUT013', 'Provide mens general grooming services', '03/11/2025'),
    ('SHBHDES001', 'Dry hair to shape', '03/11/2025'),
    ('SHBHIND001', 'Maintain and organise tools, equipment and work areas', '22/04/2025'),
    ('SHBHIND002', 'Research and use hairdressing industry information', '09/02/2026'),
    ('SHBHIND003', 'Develop and expand a client base', '22/10/2025'),
    ('SHBHTRI001', 'Identify and treat hair and scalp conditions', '05/08/2025'),
    ('SHBXCCS001', 'Conduct salon financial transactions', '16/03/2025'),
    ('SHBXCCS002', 'Provide salon services to clients', '05/08/2025'),
    ('SHBXIND001', 'Comply with organisational requirements', '03/11/2025'),
    ('SHBXIND002', 'Communicate as part of a salon team', '05/08/2025'),
    ('SHBXWHS001', 'Apply safe hygiene, health and work practices', '25/02/2025'),
    ('SIRRINV001', 'Receive and handle retail stock', '10/04/2025'),
    ('SIRRMER001', 'Produce visual merchandise displays', '13/08/2025'),
    ('SIRXSLS001', 'Sell to the retail customer', '12/01/2026'),
]
for code, name, date in units:
    row = table.add_row().cells
    row[0].text = code
    row[1].text = name
    row[2].text = date

doc.add_paragraph('')

# ===== SECTION 2 =====
doc.add_heading('2. Master of Accounting', level=1)

for label, val in [
    ('Institution:', 'Kaplan Business School'),
    ('Date of Completion:', '29 November 2024'),
    ('Certificate Number:', '24645'),
    ('Cumulative GPA:', '4.62 / 7.0'),
    ('CRICOS Provider:', '02426B'),
    ('TEQSA Provider:', 'PRV12094'),
    ('Principal Executive Officer:', 'Rob Regan'),
]:
    p = doc.add_paragraph()
    r = p.add_run(label + ' ')
    r.bold = True
    p.add_run(val)

doc.add_heading('Academic Transcript', level=2)

t2 = doc.add_table(rows=1, cols=7)
t2.style = 'Light Grid Accent 1'
for i, h in enumerate(['Year', 'Period', 'Subject', 'Subject Name', 'Mark', 'Grade', 'CP']):
    t2.rows[0].cells[i].text = h
    t2.rows[0].cells[i].paragraphs[0].runs[0].bold = True

transcript = [
    ('2024', 'T2', 'ACCM4400', 'Auditing and Assurance', '70', 'C', '4'),
    ('2024', 'T2', 'ACCM6000', 'Capstone: Accounting and Governance', '61', 'P', '4'),
    ('2024', 'T2', 'FINM4000', 'Finance', '87', 'HD', '4'),
    ('2024', 'T1', 'ACCM4100', 'Management Accounting', '59', 'P', '4'),
    ('2024', 'T1', 'ACCM6000', 'Capstone: Accounting and Governance', '45', 'F', '0'),
    ('2024', 'T1', 'DATA4500', 'Social Media Analytics', '67', 'C', '4'),
    ('2023', 'T3', 'CLWM4100', 'Taxation Law', '53', 'P', '4'),
    ('2023', 'T3', 'ECOM4000', 'Economics', '58', 'P', '4'),
    ('2023', 'T2', 'ACCM4300', 'Financial Reporting', '77', 'D', '4'),
    ('2023', 'T2', 'DATA4900', 'Innovation and Creativity in Business Analytics', '68', 'C', '4'),
    ('2023', 'T2', 'FINM4100', 'Analytics in Accounting, Finance and Economics', '78', 'D', '4'),
    ('2023', 'T1', 'ACCM4200', 'Advanced Financial Accounting', '50', 'P', '4'),
    ('2023', 'T1', 'CLWM4000', 'Business and Corporations Law', '66', 'C', '4'),
    ('2023', 'T1', 'STAM4000', 'Quantitative Methods', '63', 'P', '4'),
    ('2022', 'T3', 'ACCM4000', 'Financial Accounting', '71', 'C', '4'),
    ('2022', 'T3', 'CISM4000', 'Information Systems in Accounting', '54', 'P', '4'),
    ('2022', 'T3', 'DATA4000', 'Introduction to Business Analytics', '65', 'C', '4'),
]
for row_data in transcript:
    row = t2.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph('')

# ===== SECTION 3 =====
doc.add_heading('3. Grading System (Higher Education)', level=1)

t3 = doc.add_table(rows=6, cols=4)
t3.style = 'Light Grid Accent 1'
grading = [
    ('Grade', 'Mark (%)', 'GPA', 'Description'),
    ('High Distinction (HD)', '85-100', '7.0', 'Excellent'),
    ('Distinction (D)', '75-84', '6.0', 'Very Good'),
    ('Credit (C)', '65-74', '5.0', 'Good'),
    ('Pass (P)', '50-64', '4.0', 'Satisfactory'),
    ('Fail (F)', '0-49', '1.5', 'Unsatisfactory'),
]
for i, row_data in enumerate(grading):
    for j, val in enumerate(row_data):
        t3.cell(i, j).text = val
        if i == 0:
            t3.cell(i, j).paragraphs[0].runs[0].bold = True

doc.add_paragraph('')

# Summary
doc.add_heading('Summary', level=1)
bullets = [
    'Certificate III in Barbering (SHB30516) — Completed March 2026 — The Barber Academy / Propel Education (RTO #41209)',
    'Master of Accounting — Completed November 2024 — Kaplan Business School (CRICOS 02426B, TEQSA PRV12094)',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

output = r'C:\Users\Dewa(Interlace)\Documents\Interlace Code\WebScraping-CourseCricos\Onshore_Academics_Anil_Thapa.docx'
doc.save(output)
print(f'DOCX saved: {output}')
